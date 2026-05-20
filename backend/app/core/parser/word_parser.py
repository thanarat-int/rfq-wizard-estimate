from docx import Document
import io


def extract_text_from_word(file_content: bytes) -> str:
    """Extract text content from a Word document."""
    doc = Document(io.BytesIO(file_content))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract tables
    for table in doc.tables:
        text_parts.append("--- Table ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            text_parts.append(row_text)

    return "\n".join(text_parts)
